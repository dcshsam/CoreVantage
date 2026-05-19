import io


def parse_document(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return _parse_pdf(uploaded_file)
    elif name.endswith(".docx"):
        return _parse_docx(uploaded_file)
    elif name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {name}")


def _parse_pdf(file) -> str:
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        result = "\n".join(parts)
        if result.strip():
            return result
    except Exception:
        pass

    # Fallback to PyPDF2
    import PyPDF2
    file.seek(0)
    reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)


def _parse_docx(file) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file.read()))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    all_text = paragraphs + table_rows
    return "\n".join(all_text)
