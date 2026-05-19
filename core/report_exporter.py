"""
CodeVantage Report Exporter
Generates professional PDF, Word, and Excel reports from analysis results.
"""

from __future__ import annotations
import io
from datetime import datetime
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from core.clean_code_engine import CleanCoreResult
    from core.s4_migration_engine import MigrationResult


# ── Word (.docx) ──────────────────────────────────────────────────────────────

def export_clean_core_word(result: "CleanCoreResult") -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _add_cv_header(doc, "SAP Clean Core Analysis Report", result.program_name)

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"Program: {result.program_name} | "
        f"Lines: {result.total_lines:,} | "
        f"Violations: {result.total_violations} | "
        f"Clean Core Level: {result.clean_core_level}"
    )
    if result.llm_analysis:
        doc.add_paragraph(result.llm_analysis.split("\n\n")[0])

    # Violation summary table
    doc.add_heading("Violation Summary", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Rule ID", "Severity", "Line", "Issue", "Remediation"]):
        hdr[i].text = h
    for v in result.violations[:100]:
        row = table.add_row().cells
        row[0].text = v.rule_id
        row[1].text = v.severity
        row[2].text = str(v.line_number)
        row[3].text = v.rule.name[:60]
        row[4].text = v.remediation[:80]

    # LLM analysis
    if result.llm_analysis:
        doc.add_heading("AI-Powered Analysis", level=1)
        for para in result.llm_analysis.split("\n\n"):
            if para.startswith("##"):
                doc.add_heading(para.lstrip("#").strip(), level=2)
            elif para.strip():
                doc.add_paragraph(para.strip())

    # Remediated code
    if result.remediated_code:
        doc.add_heading("Remediated Code", level=1)
        p = doc.add_paragraph()
        run = p.add_run(result.remediated_code[:5000])
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    _add_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_s4_migration_word(result: "MigrationResult") -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _add_cv_header(doc, "S/4HANA Migration Analysis Report", result.program_name)

    doc.add_heading("Migration Readiness", level=1)
    doc.add_paragraph(
        f"Program: {result.program_name} | "
        f"Readiness Score: {result.readiness_score}/100 | "
        f"Risk: {result.risk_level} | "
        f"Effort: {result.effort_days} days | "
        f"Approach: {result.approach}"
    )

    if result.s4_violations:
        doc.add_heading("Migration Issues", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Rule ID", "Severity", "Line", "Issue"]):
            hdr[i].text = h
        for v in result.s4_violations[:80]:
            row = table.add_row().cells
            row[0].text = v.rule_id
            row[1].text = v.severity
            row[2].text = str(v.line_number)
            row[3].text = v.rule.name[:80]

    if result.migration_plan:
        doc.add_heading("Migration Plan", level=1)
        for para in result.migration_plan.split("\n\n"):
            if para.startswith("##"):
                doc.add_heading(para.lstrip("#").strip(), level=2)
            elif para.strip():
                doc.add_paragraph(para.strip())

    _add_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Excel (.xlsx) ─────────────────────────────────────────────────────────────

def export_violations_excel(violations, program_name: str = "Analysis") -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb  = openpyxl.Workbook()
    ws  = wb.active
    ws.title = "Violations"

    COLORS = {
        "CRITICAL": "FFBB0000",
        "HIGH":     "FFE9730C",
        "MEDIUM":   "FFC87400",
        "LOW":      "FF0070F2",
        "INFO":     "FF6D6D6D",
    }

    headers = ["Rule ID", "Category", "Severity", "Line #", "Rule Name",
               "Code Snippet", "Description", "Remediation"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFFFF")
        cell.fill = PatternFill("solid", fgColor="FF003F7F")
        cell.alignment = Alignment(wrap_text=True)

    for row_idx, v in enumerate(violations, 2):
        data = [v.rule_id, v.category, v.severity, v.line_number,
                v.rule.name, v.line_content[:100], v.description[:200], v.remediation[:200]]
        for col, val in enumerate(data, 1):
            cell = ws.cell(row=row_idx, column=col, value=val)
            if col == 3:  # Severity
                color = COLORS.get(v.severity, "FF6D6D6D")
                cell.fill = PatternFill("solid", fgColor=color)
                cell.font = Font(color="FFFFFFFF", bold=True)
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = 20
    ws.column_dimensions["F"].width = 40
    ws.column_dimensions["H"].width = 50

    # Summary sheet
    ws2 = wb.create_sheet("Summary")
    ws2["A1"] = "CodeVantage Analysis Summary"
    ws2["A1"].font = Font(bold=True, size=14)
    ws2["A3"] = "Program"
    ws2["B3"] = program_name
    ws2["A4"] = "Total Violations"
    ws2["B4"] = len(violations)
    ws2["A5"] = "Generated"
    ws2["B5"] = datetime.now().strftime("%Y-%m-%d %H:%M")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────────────

def export_pdf(title: str, content: str, program_name: str = "") -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles  = getSampleStyleSheet()
    story   = []
    SAP_BLUE = colors.HexColor("#0A6ED1")

    title_style = ParagraphStyle("cv_title", parent=styles["Title"],
                                  textColor=SAP_BLUE, fontSize=18, spaceAfter=6)
    h1_style    = ParagraphStyle("cv_h1", parent=styles["Heading1"],
                                  textColor=SAP_BLUE, fontSize=13)
    body_style  = ParagraphStyle("cv_body", parent=styles["Normal"],
                                  fontSize=9, leading=13)

    story.append(Paragraph("⚡ CodeVantage", title_style))
    story.append(Paragraph(title, h1_style))
    if program_name:
        story.append(Paragraph(f"Program: {program_name}", body_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}", body_style))
    story.append(Spacer(1, 0.5*cm))

    for line in content.split("\n"):
        if line.startswith("##"):
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph(line.lstrip("#").strip(), h1_style))
        elif line.strip():
            story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;"), body_style))

    doc.build(story)
    return buf.getvalue()


# ── Helpers ───────────────────────────────────────────────────────────────────

def _add_cv_header(doc, title: str, program_name: str) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("⚡ CodeVantage — Enterprise ABAP Intelligence Platform")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0A, 0x6E, 0xD1)

    doc.add_heading(title, level=1)
    doc.add_paragraph(
        f"Program: {program_name}   |   "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}"
    )
    doc.add_paragraph("─" * 80)


def _add_footer(doc) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph("─" * 80)
    p = doc.add_paragraph("Generated by CodeVantage — Enterprise ABAP Intelligence Platform")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
